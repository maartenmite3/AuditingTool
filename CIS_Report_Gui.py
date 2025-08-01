import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
from datetime import datetime
import hashlib
import os
import ttkbootstrap as tb
from ttkbootstrap.constants import *
from tkinter import filedialog

def normalize_text(text):
    return ' '.join(text.lower().split())

def extract_paragraphs_clean(soup_section):
    output = []
    seen_hashes = set()
    for tag in soup_section.find_all(["p", "li"], recursive=True):
        if tag.find(["p", "li"]):  # skip nested <p> or <li>
            continue
        text = tag.get_text(" ", strip=True)
        if not text:
            continue
        norm = normalize_text(text)
        hash_val = hashlib.md5(norm.encode()).hexdigest()
        if hash_val not in seen_hashes:
            seen_hashes.add(hash_val)
            output.append(f"- {text}" if tag.name == "li" else text)
    return output

def extract_clean_text(html):
    soup = BeautifulSoup(html, "html.parser")
    lines = []
    seen = set()
    for elem in soup.find_all(["p", "li"]):
        text = elem.get_text(" ", strip=True)
        norm = normalize_text(text)
        if norm and norm not in seen:
            seen.add(norm)
            lines.append(f"- {text}" if elem.name == "li" else text)
    return "\n".join(lines).strip()

def extract_impact_section(html):
    soup = BeautifulSoup(html, "html.parser")
    impact_header = soup.find("p", class_="bold")
    if impact_header and "impact" in impact_header.text.lower():
        parts = str(soup).split(str(impact_header), 1)
        impact = BeautifulSoup(parts[1], "html.parser")
        lines = extract_paragraphs_clean(impact)
        unique_lines = sorted(set(l.strip() for l in lines if l.strip()))
        return unique_lines
    return []

def extract_remediation_section(html):
    soup = BeautifulSoup(html, "html.parser")
    impact_header = soup.find("p", class_="bold")
    if impact_header and "impact" in impact_header.text.lower():
        parts = str(soup).split(str(impact_header), 1)
        remediation_html = BeautifulSoup(parts[0], "html.parser")
    else:
        remediation_html = soup

    result = []
    previous_text = ""
    for tag in remediation_html.find_all(["p", "li"], recursive=True):
        text = tag.get_text(" ", strip=True)
        if not text or text == previous_text:
            continue
        previous_text = text

        if tag.find("code"):
            intro = ""
            code = ""
            for content in tag.contents:
                if hasattr(content, "name") and content.name == "code":
                    code = content.get_text(" ", strip=True)
                else:
                    intro += content.get_text(" ", strip=True) if hasattr(content, 'get_text') else str(content)
            intro = intro.strip().rstrip(":")
            if intro:
                result.append(("text", intro + ":"))
            if code:
                result.append(("code", code))
        else:
            if text not in [val for (_, val) in result]:
                result.append(("text", text))

    cleaned_result = []
    skip_next = False
    for i, (kind, val) in enumerate(result):
        if skip_next:
            skip_next = False
            continue
        if kind == "code" and i + 1 < len(result):
            next_kind, next_val = result[i + 1]
            if next_kind == "text" and next_val.lower().startswith("to establish"):
                cleaned_result.append((kind, val))
                skip_next = True
            else:
                cleaned_result.append((kind, val))
        else:
            cleaned_result.append((kind, val))
    return cleaned_result

def add_cover_page(doc, computer_name):
    for _ in range(12):
        doc.add_paragraph("")

    title = doc.add_paragraph("CIS COMPLIANCE REPORT")
    title.alignment = 1
    title_run = title.runs[0]
    title_run.font.size = Pt(28)
    title_run.bold = True

    doc.add_paragraph("")
    subtitle = doc.add_paragraph("System Overview and Findings")
    subtitle.alignment = 1
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(16)

    doc.add_paragraph("")
    doc.add_paragraph("")

    date_text = f"Generated on: {datetime.now().strftime('%d-%m-%Y')}"
    computer_text = f"Target System: {computer_name}"

    date_para = doc.add_paragraph()
    date_para.alignment = 1
    date_run = date_para.add_run(date_text)
    date_run.font.size = Pt(18)

    comp_para = doc.add_paragraph()
    comp_para.alignment = 1
    comp_run = comp_para.add_run(computer_text)
    comp_run.font.size = Pt(18)

    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_page_break()


def add_table_of_contents(doc, toc_entries):
    doc.add_heading("Table of Contents", level=1)
    for entry in toc_entries:
        doc.add_paragraph(entry)
    doc.add_page_break()

def generate_cis_report(df, output_path, computer_name="", include_cover=True, include_toc=True):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    if include_cover:
        add_cover_page(doc, computer_name)

    if include_toc:
        toc_entries = [
            f"{row['ConfigurationId']} – {row['ConfigurationName']}"
            for _, row in df.iterrows()
        ]
        add_table_of_contents(doc, toc_entries)

    doc.add_heading("CIS Compliance Findings", level=1)
    doc.add_paragraph("")

    for idx, (_, row) in enumerate(df.iterrows()):
        if idx > 0:
            doc.add_page_break()

        config_line = f"{row['ConfigurationId']} – {row['ConfigurationName']}"
        para = doc.add_paragraph()
        para.add_run(config_line).bold = True

        rationale_text = extract_clean_text(row.get('ConfigurationRationale', ''))
        doc.add_paragraph("Rationale:", style='Heading 2')
        for line in rationale_text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())

        raw_html = row.get('RemediationOptions', '')
        remediation_parts = extract_remediation_section(raw_html)
        doc.add_paragraph("Remediation:", style='Heading 2')
        for kind, content in remediation_parts:
            if kind == "text":
                doc.add_paragraph(content)
            elif kind == "code":
                p = doc.add_paragraph()
                run = p.add_run(content)
                run.font.name = "Courier New"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")

        impact_lines = extract_impact_section(raw_html)
        if impact_lines:
            doc.add_paragraph("Impact:", style='Heading 2')
            for line in impact_lines:
                doc.add_paragraph(line)

        doc.add_paragraph("Comply or Explain:", style='Heading 2')
        doc.add_paragraph("Explain: ........................................................................................................")
        doc.add_paragraph("☐ Exception / Exemption")
        doc.add_paragraph("─" * 78)

    doc.save(output_path)
    print(f"✅ Report saved: {output_path}")

def run_gui():
    def generate_report():
        file_path = file_entry.get()
        device_name = device_entry.get().strip().replace(" ", "_")
        computer_name = computer_entry.get().strip()
        include_cover = bool(include_cover_var.get())
        include_toc = bool(include_toc_var.get())

        if not os.path.exists(file_path) or not device_name:
            status_label.config(text="⚠️ Provide valid CSV file and device name", bootstyle="danger")
            return
        try:
            df = pd.read_csv(file_path)
            output_path = f"cis_compliance_report_{device_name}.docx"
            generate_cis_report(
                df, output_path,
                computer_name=computer_name,
                include_cover=include_cover,
                include_toc=include_toc
            )
            status_label.config(text=f"✅ Report saved: {output_path}", bootstyle="success")
        except Exception as e:
            status_label.config(text=f"❌ Error: {e}", bootstyle="danger")

    def browse_file():
        filename = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if filename:
            file_entry.delete(0, "end")
            file_entry.insert(0, filename)

    app = tb.Window(themename="flatly")
    app.title("CIS Compliance Report Generator")
    app.geometry("600x420")

    include_cover_var = tb.IntVar(value=1)
    include_toc_var = tb.IntVar(value=1)

    tb.Label(app, text="Select CIS CSV File:").pack(pady=(10, 2))
    file_frame = tb.Frame(app)
    file_frame.pack(fill=X, padx=10)
    file_entry = tb.Entry(file_frame, width=50)
    file_entry.pack(side=LEFT, fill=X, expand=True)
    tb.Button(file_frame, text="Browse", command=browse_file).pack(side=RIGHT, padx=5)

    tb.Label(app, text="Device Name (used in output file):").pack(pady=(10, 2))
    device_entry = tb.Entry(app, width=50)
    device_entry.pack(padx=10)

    tb.Label(app, text="Computer/Host Name (for cover page):").pack(pady=(10, 2))
    computer_entry = tb.Entry(app, width=50)
    computer_entry.pack(padx=10)

    tb.Checkbutton(app, text="Include Cover Page", variable=include_cover_var).pack(anchor="w", padx=10, pady=(10, 2))
    tb.Checkbutton(app, text="Include Table of Contents", variable=include_toc_var).pack(anchor="w", padx=10)

    tb.Button(app, text="Generate Report", command=generate_report, bootstyle=PRIMARY).pack(pady=20)
    status_label = tb.Label(app, text="", font=("Calibri", 10))
    status_label.pack()

    app.mainloop()

if __name__ == "__main__":
    run_gui()
