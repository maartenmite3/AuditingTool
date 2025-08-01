
import csv
import requests
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from time import sleep
from tkinter import Tk, filedialog, messagebox
from ttkbootstrap import Style
from ttkbootstrap.widgets import Entry, Button, Label, Frame

def load_microsoft_data(csv_path):
    df = pd.read_csv(csv_path, skiprows=1)
    df.columns = [col.strip() for col in df.columns]
    df.columns.values[0:4] = ['CVE_ID', 'Severity', 'CVSS v3', 'Description']
    df = df[['CVE_ID', 'Severity', 'CVSS v3', df.columns[-1]]].dropna(subset=["CVE_ID"])
    df = df.rename(columns={df.columns[-1]: "Description"})
    return df

def get_api_key_from_file(file_path='nvd_api_key.txt'):
    try:
        with open(file_path, 'r') as f:
            return f.read().strip()
    except FileNotFoundError:
        raise RuntimeError(f"API key file not found: {file_path}")

def get_nvd_data(cve_id, api_key):
    try:
        headers = {'apiKey': api_key}
        response = requests.get(f'https://services.nvd.nist.gov/rest/json/cves/2.0?cveId={cve_id}', headers=headers)
        response.raise_for_status()
        data = response.json()
        vuln = data['vulnerabilities'][0]['cve']
        title = vuln['descriptions'][0]['value']
        score = 'N/A'
        severity = 'Unknown'
        metrics = vuln.get('metrics', {})

        if 'cvssMetricV31' in metrics:
            cvss = metrics['cvssMetricV31'][0]['cvssData']
        elif 'cvssMetricV30' in metrics:
            cvss = metrics['cvssMetricV30'][0]['cvssData']
        else:
            cvss = None

        if cvss:
            score = cvss.get('baseScore', 'N/A')
            try:
                s = float(score)
                if s == 0.0:
                    severity = 'None'
                elif s < 4.0:
                    severity = 'Low'
                elif s < 7.0:
                    severity = 'Medium'
                elif s < 9.0:
                    severity = 'High'
                else:
                    severity = 'Critical'
            except:
                pass

        return {
            'title': title,
            'severity': severity,
            'score': score,
            'link': f'https://nvd.nist.gov/vuln/detail/{cve_id}'
        }

    except Exception as e:
        print(f"[ERROR] NVD fetch failed for {cve_id}: {e}")
        return {
            'title': '[Failed to fetch NVD title]',
            'severity': 'Unknown',
            'score': 'N/A',
            'link': f'https://nvd.nist.gov/vuln/detail/{cve_id}'
        }

def write_intro_page(doc):
    doc.add_heading("CVE Severity Report (Microsoft + NVD)", level=1)
    doc.add_paragraph(
        "This report compares vulnerability severity ratings from two sources:\n"
        "- Microsoft: contextual severity based on asset exposure, exploitability, and threat landscape.\n"
        "- NVD (National Vulnerability Database): technical severity from a standardized CVSS scoring system.\n\n"
        "Both perspectives are valuable — Microsoft focuses on actionable risk, while NVD reflects inherent technical impact."
    )

def insert_cve_section(doc, cve_id, ms_severity, ms_cvss, ms_desc, nvd_info, is_first=False):
    if not is_first:
        doc.add_page_break()
    para = doc.add_paragraph()
    run = para.add_run(f"CVE ID: {cve_id} – {nvd_info['title']}")
    run.bold = True
    para.paragraph_format.left_indent = Inches(0)

    doc.add_paragraph("Microsoft Severity:", style='Heading 2')
    doc.add_paragraph(f"{ms_severity} (CVSS {ms_cvss})")

    doc.add_paragraph("NVD Severity:", style='Heading 2')
    doc.add_paragraph(f"{nvd_info['severity']} (CVSS {nvd_info['score']})")

    doc.add_paragraph("Risk Description:", style='Heading 2')
    doc.add_paragraph(ms_desc)

    doc.add_paragraph("Mitigation Advice:", style='Heading 2')
    doc.add_paragraph("Refer to vendor patches, Microsoft Defender TVM guidance, or mitigation techniques.")

    doc.add_paragraph("References:", style='Heading 2')
    doc.add_paragraph(f"NVD: {nvd_info['link']}")

    doc.add_paragraph("Comply or Explain:", style='Heading 2')
    doc.add_paragraph("Explain: ........................................................................................................")
    doc.add_paragraph("☐ Exception / Exemption")
    doc.add_paragraph("─" * 78)

def generate_report(csv_path, template_path, delay, output_path):
    api_key = get_api_key_from_file()
    df = load_microsoft_data(csv_path)
    doc = Document(template_path)
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)
    write_intro_page(doc)

    for idx, row in df.iterrows():
        cve_id = str(row['CVE_ID']).strip()
        if not cve_id.startswith("CVE-"):
            print(f"Skipping invalid CVE ID: {cve_id}")
            continue
        ms_sev = row['Severity']
        ms_score = row['CVSS v3']
        ms_desc = str(row['Description'])

        print(f"Processing {cve_id}...")
        nvd_info = get_nvd_data(cve_id, api_key)
        insert_cve_section(doc, cve_id, ms_sev, ms_score, ms_desc, nvd_info, is_first=(idx == 0))
        sleep(delay)

    doc.save(output_path)
    print(f"Report saved to {output_path}")

def launch_gui():
    def browse_csv():
        path = filedialog.askopenfilename(filetypes=[("CSV Files", "*.csv")])
        if path:
            csv_entry.delete(0, 'end')
            csv_entry.insert(0, path)

    def browse_template():
        path = filedialog.askopenfilename(filetypes=[("Word Documents", "*.docx")])
        if path:
            template_entry.delete(0, 'end')
            template_entry.insert(0, path)

    def run_script():
        csv_path = csv_entry.get().strip()
        template_path = template_entry.get().strip()
        delay = float(delay_entry.get().strip())
        device_name = device_entry.get().strip()
        output_path = f"{device_name}_vuln_report.docx"

        if not all([csv_path, template_path, device_name]):
            messagebox.showerror("Input Error", "Please fill in all fields.")
            return

        try:
            generate_report(csv_path, template_path, delay, output_path)
            messagebox.showinfo("Done", f"Report saved as {output_path}")
        except Exception as e:
            messagebox.showerror("Error", str(e))

    root = Tk()
    root.title("CVE Word Report Generator")
    style = Style(theme="superhero")

    frame = Frame(root, padding=10)
    frame.pack(fill='both', expand=True)

    Label(frame, text="CSV File:").grid(row=0, column=0, sticky='w')
    csv_entry = Entry(frame, width=50)
    csv_entry.grid(row=0, column=1)
    Button(frame, text="Browse", command=browse_csv).grid(row=0, column=2)

    Label(frame, text="Word Template:").grid(row=1, column=0, sticky='w')
    template_entry = Entry(frame, width=50)
    template_entry.grid(row=1, column=1)
    Button(frame, text="Browse", command=browse_template).grid(row=1, column=2)

    Label(frame, text="API Delay (s):").grid(row=2, column=0, sticky='w')
    delay_entry = Entry(frame, width=10)
    delay_entry.insert(0, "1.5")
    delay_entry.grid(row=2, column=1, sticky='w')

    Label(frame, text="Device Name (for output):").grid(row=3, column=0, sticky='w')
    device_entry = Entry(frame, width=30)
    device_entry.grid(row=3, column=1, columnspan=2)

    Button(frame, text="Generate Report", bootstyle="success", command=run_script).grid(row=4, column=1, pady=10)

    root.mainloop()

if __name__ == "__main__":
    launch_gui()
